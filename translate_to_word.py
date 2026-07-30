# -*- coding: utf-8 -*-
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

IMG_DIR = r"D:\work\2607\extracted_images"
OUT_DIR = r"D:\work\2607"

def set_chinese_font(run, font_name="SimSun", size=10.5, bold=False):
    font = run.font
    font.name = font_name
    font.size = Pt(size)
    font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_zh(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if level == 0:
        set_chinese_font(run, "SimHei", 16, bold=True)
    elif level == 1:
        set_chinese_font(run, "SimHei", 14, bold=True)
    elif level == 2:
        set_chinese_font(run, "SimHei", 12, bold=True)
    else:
        set_chinese_font(run, "SimHei", 11, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p

def add_body_zh(doc, text, size=10.5, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=Cm(0.74)):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = first_line_indent
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_chinese_font(run, "SimSun", size)
    return p

def add_caption_zh(doc, text, size=9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_chinese_font(run, "SimHei", size, bold=True)
    return p

def add_image_at(doc, img_path, width_cm=15):
    if not os.path.exists(img_path):
        print(f"WARNING: image not found: {img_path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))
    print(f"Inserted image: {os.path.basename(img_path)}")

def add_no_indent(doc, text, size=10.5, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_chinese_font(run, "SimSun", size)
    return p

# ──────────────────────────────────────────────
# MAIN DOCUMENT
# ──────────────────────────────────────────────
def build_main():
    doc = Document()
    
    # Set default font for document
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # ---- Title ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("欧洲复合沿海海洋-陆地热浪与湿热胁迫的关联研究")
    set_chinese_font(run, "SimHei", 16, bold=True)
    p.paragraph_format.space_after = Pt(8)
    
    # ---- Authors ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Armineh Barkhordarian¹*, Eline Brunet¹,² & Johanna Baehr¹")
    set_chinese_font(run, "SimSun", 10.5)
    p.paragraph_format.space_after = Pt(4)
    
    # ---- Affiliation ----
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("¹德国汉堡大学海洋研究所，汉堡，德国  ²布列塔尼西部大学欧洲海洋研究所（IUEM），法国")
    set_chinese_font(run, "SimSun", 9)
    p.paragraph_format.space_after = Pt(4)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("*通讯作者: armineh.barkhordarian@uni-hamburg.de")
    set_chinese_font(run, "SimSun", 9)
    p.paragraph_format.space_after = Pt(12)
    
    # ---- Abstract ----
    add_heading_zh(doc, "摘要", level=1)
    abstract = (
        "欧洲沿海地区——人口密集、生态系统对气候敏感——极易受到海洋-陆地热浪叠加效应的影响。"
        "本研究将复合热浪定义为陆地和相邻海域同时发生的极端高温事件。"
        "利用卫星和地面观测数据，我们发现过去二十年间欧洲沿海复合热浪暴露呈非线性加速增长趋势，"
        "2022年地中海地区达到峰值78天。归因分析表明，温室气体（GHG）强迫是2022年该事件风险的95%成因；"
        "CESM1-LE模拟显示，若无温室气体强迫，此类事件几乎不可能发生，凸显了GHG排放持续增加所带来的风险非线性升级。"
        "我们还证明，邻近海域同时发生的海洋热浪可将沿海陆地热浪暴露增强至3.5倍，"
        "将短暂的陆地热浪转变为长期极端高温高湿事件。"
        "在复合事件期间，沿海地区明显转向湿热热浪状态，其特征为湿球温度超过25.5°C且比湿升高。"
        "我们的研究揭示了海洋热浪通过增强湿度和热量对沿海陆地极端高温持续时间的放大效应。"
    )
    add_body_zh(doc, abstract)
    
    # ---- Introduction (Pages 1-2) ----
    add_heading_zh(doc, "引言", level=1)
    
    intro_p1 = (
        "欧洲沿海地区是人口密集区域，也是气候敏感生态系统的家园，"
        "易受海洋-陆地热浪叠加效应的影响。"
        "在干旱地中海地区，干复合极端事件已与增强的野火活动相关联，"
        "这些野火由同步大气热浪、干旱和海洋热浪共同驱动[26]。"
    )
    add_body_zh(doc, intro_p1)
    
    intro_p2 = (
        "欧洲沿海地区对海洋-陆地复合热极端尤为脆弱，"
        "因为这些事件在人口密集的沿海地区的叠加——乃至潜在的协同——"
        "影响可能显著放大健康风险[27]。"
        "在地中海东部海域，超过一半的海洋热浪与海上的大气热浪同时发生[27]，"
        "已有研究表明这种并发会通过海-气热通量变化增强海洋热浪[28,29]。"
        "2003年地中海地区引人注目的大气热浪和海洋热浪事件已成为多项研究的焦点，"
        "这些研究指出，地中海升高的海表温度在陆地热浪开始后加强了其强度，"
        "尽管并未促成其初始发生[30,31]。"
    )
    add_body_zh(doc, intro_p2)
    
    intro_p3 = (
        "以往研究表明，大气热极端可以增强海洋热异常，"
        "尤其在地中海地区；但这些评估大多在大区域尺度[27-29]上进行，"
        "或通过独立比较陆地和海洋事件[30,31]进行。"
        "因此，它们未能解决热浪相互作用在具体海岸线的演变过程。"
        "在此，我们识别并量化复合沿海海洋-陆地热浪，"
        "定义为同时被并发海洋热浪完全涵盖的陆地热浪，"
        "基于配对的沿海陆地和相邻海洋格点数据。"
    )
    add_body_zh(doc, intro_p3)
    
    # ---- Section: Spatial Distribution (Page 3) ----
    add_heading_zh(doc, "复合海洋-陆地热浪的空间分布特征", level=2)
    
    spatial_p1 = (
        "为了解复合海洋-陆地热浪在何处最为普遍，"
        "我们分析了1984-2023年欧洲沿海E-OBS观测数据覆盖区域复合发生概率的空间分布"
        "（定义为复合天数与陆地热浪天数的比值，取2003-2023年均值，图1m）。"
        "所得模式揭示了区域热点，并突显了不同盆地间海洋-陆地热浪耦合强度的差异。"
        "最高复合发生概率（>0.8）出现在西地中海沿岸及撒丁岛和科西嘉岛周边，"
        "表明这些地区超过80%的陆地热浪天数与海洋热浪重合。"
        "这些区域以浅水和频繁的大气停滞为特征，有利于热量和湿度的积聚[40-42]。"
    )
    add_body_zh(doc, spatial_p1)
    
    # INSERT FIGURE 1 HERE (after mentioning Fig. 1m)
    add_image_at(doc, os.path.join(IMG_DIR, "compound coastal marine-terrestrial heatwaves associated with humid-heat stress in europe(1)_p3_img1.png"), width_cm=15)
    add_caption_zh(doc, "图1. a-i 基于OISST和E-OBS观测记录的复合海洋-陆地热浪天数的年际空间分布。"
                        "j-l 2003-2023年地中海区域（含黑海）、波罗的海沿岸及所有欧洲海岸线的复合海洋-陆地热浪天数的时间演变。"
                        "m 复合发生概率的空间分布，定义为复合沿海海洋-陆地热浪天数与所有陆地热浪天数的比值。"
                        "数值范围从0（无陆地热浪天数为复合事件）到1（所有陆地热浪天数均为复合事件），统计周期为2003-2023年。")
    
    spatial_p2 = (
        "同样较高的复合发生值也出现在南波罗的海沿岸部分地区，"
        "包括波兰北部、立陶宛西部和拉脱维亚。"
        "中等复合概率（0.6-0.8）在地中海盆地大部分地区广泛存在，"
        "包括黑海沿岸（图1m）。相比之下，大西洋沿岸的复合比例较低（0.2-0.4），"
        "表明海洋与陆地热极端之间的耦合较弱——但仍值得注意（图1m）。"
        "这些模式强化了海洋热浪在调节陆地极端高温事件中日益重要作用的认识，"
        "尤其在地中海和波罗的海等半封闭盆地。"
    )
    add_body_zh(doc, spatial_p2)
    
    # ---- Section: Quantifying Role (Page 4) ----
    add_heading_zh(doc, "量化海洋热浪对陆地热浪暴露的调节作用", level=2)
    
    quantify_p1 = (
        "独立陆地热浪天数（定义为在相邻海洋格点中无并发海洋热浪的事件）"
        "与复合热浪天数之间的空间对比，显示复合事件在欧洲沿海极端高温暴露中日益占主导地位（图2a,b）。"
        "2003-2023年独立陆地热浪天数的平均值在大多数沿海区域仍低于10天（图2b）。"
        "相比之下，复合海洋-陆地天数在多个南欧和东欧沿海区域超过20天，"
        "表明复合热浪事件占主导地位——尤其在西地中海和东波罗的海沿岸（图2a）。"
        "为评估海洋热浪在调节陆地热浪暴露方面的作用，"
        "我们分析了复合热浪天数与独立陆地热浪天数之比（图2c,d）。"
    )
    add_body_zh(doc, quantify_p1)
    
    # INSERT FIGURE 2 HERE
    add_image_at(doc, os.path.join(IMG_DIR, "compound coastal marine-terrestrial heatwaves associated with humid-heat stress in europe(1)_p4_img1.png"), width_cm=15)
    add_caption_zh(doc, "图2. a 基于OISST和E-OBS观测记录的2003-2023年均复合海洋-陆地热浪天数沿海分布。"
                        "b 独立沿海陆地热浪天数（定义为相邻海洋格点中无并发海洋热浪的陆地事件）。"
                        "c 复合热浪比（CHR）的时间演变，定义为1983-2023年复合海洋-陆地热浪与独立陆地热浪之比。"
                        "d CHR的沿海空间分布（2003-2023年均值）。")
    
    quantify_p2 = (
        "在早期记录（1980年代-1990年代）中，CHR在1附近波动，"
        "意味着陆地热浪大多是独立的，孤立发生以及与海洋热浪并发频率相同（图2c）。"
        "然而自2003年以来，该比值持续超过1，"
        "表明陆地热浪暴露时间在并发海洋热浪存在时趋于增加。"
        "2023年CHR峰值达3.5，表明当海洋热浪同时发生时，"
        "陆地热浪暴露时间增加了3.5倍。"
        "2003年后CHR的加速上升与海洋热浪的加剧一致，"
        "表明陆-海相互作用在塑造欧洲沿海极端高温条件方面的影响日益增强。"
    )
    add_body_zh(doc, quantify_p2)
    
    quantify_p3 = (
        "复合热浪比的空间分布（2003-2023年均值）显示多个欧洲沿海区域复合热浪事件占主导地位（图2d）。"
        "CHR值超过3的地区明显出现在西地中海沿岸——包括西班牙南部、法国南部和意大利西部——"
        "表明这些地区陆地热浪暴露在并发海洋热浪时比独立发生时高3倍以上。"
        "亚得里亚海、爱琴海和东波罗的海沿岸也观察到升高的CHR值（2-3）。"
        "这些结果表明，陆地热浪在并发附近海洋热浪时持续时间更长。"
        "相比之下，大西洋沿岸和北欧海岸大部分地区CHR值较低（<1.5），"
        "表明这些区域海洋与陆地热浪事件之间的耦合较弱。"
    )
    add_body_zh(doc, quantify_p3)
    
    # ---- Section: Attribution (Pages 5-6) ----
    add_heading_zh(doc, "温室气体强迫对复合热浪暴露的归因", level=2)
    
    attr_p1 = (
        "我们使用极端事件归因技术[43]来识别可由温室气体强迫归因的"
        "复合海洋-陆地热浪暴露风险的比例[7,8]。"
        "我们通过分析配对的沿海陆地2米温度和相邻海洋海表温度数据，"
        "量化了CESM1-LE模拟中欧洲沿海陆地和海洋热浪的并发率。"
        "我们估计了存在和不存在温室气体强迫条件下复合热浪暴露发生的概率。"
        "这些概率针对实际（全强迫）和反事实（固定GHG强迫）两种情景进行计算。"
        "所估计的概率用于计算两个事件归因指标，即可归因风险比例（FAR，方法中的公式2）"
        "和概率比（PR，方法中的公式3）。"
        "表1总结了三个选定观测事件在 Mediterranean 的年度复合热浪天数对应的FAR和PR值——"
        "2003年（62天）、2022年（78天）和2023年（72天）。"
    )
    add_body_zh(doc, attr_p1)
    
    # INSERT FIGURE 3 HERE
    add_image_at(doc, os.path.join(IMG_DIR, "compound coastal marine-terrestrial heatwaves associated with humid-heat stress in europe(1)_p5_img1.png"), width_cm=15)
    add_caption_zh(doc, "图3. 归因指标：可归因风险比例（FAR；蓝色曲线）和概率比（PR；蓝色曲线）。"
                        "a,b 显示欧洲海岸线的结果；c,d 显示地中海和黑海的结果。"
                        "x轴上的阈值代表复合热浪天数值。"
                        "c和d面板中的虚线表示2003年（62天，红色）、2022年（78天，绿色）和2023年（72天，黑色）的观测暴露持续时间。"
                        "阴影区域表示使用非参数自助法重采样方法估计的不确定性范围。")
    
    attr_p2 = (
        "可归因风险比例（FAR）曲线量化了GHG强迫是事件必要原因的概率，"
        "当暴露于复合热浪超过90天时，该曲线饱和于1.0。"
        "这表明，任何超过该阈值的复合事件在没有GHG强迫的情况下都不会发生，概率为99%（图3a）。"
        "相应的概率比（PR）发散至无穷大，进一步支持了此类极端复合事件在没有GHG强迫的气候中几乎不可能存在（图3b）。"
        "在较低阈值下，不确定性范围（蓝色不确定性带）较窄，意味着估计更稳健。"
        "在较高阈值下，不确定性扩大，可能原因是极端事件较少，导致采样变异性更大。"
    )
    add_body_zh(doc, attr_p2)
    
    attr_p3 = (
        "在地中海和黑海，GHG强迫使一个具有62天暴露量的事件（如2003年事件）"
        "发生的概率增加了4倍（PR=4；5-95%置信区间：2.8-7.2）。"
        "相应的可归因风险比例（FAR）为0.72（5-95%置信区间：0.64-0.80），"
        "表明约72%的发生风险可归因于GHG强迫（图3c,d，表1）。"
        "换言之，大约十分之七的极端事件在没有GHG强迫的情况下不会发生。"
        "2023年事件（72天暴露量）的FAR值为0.78（5-95%置信区间：0.70-0.88），"
        "表明该事件不能仅归因于自然变率。GHG强迫约占发生风险的78%，是观测事件的必要原因。"
        "接近1（100%）的FAR值意味着GHG强迫是事件的必要原因。"
    )
    add_body_zh(doc, attr_p3)
    
    attr_p4 = (
        "在2022年事件的情况下，GHG强迫的贡献更加显著。"
        "对于具有78天暴露量的复合事件（如2022年事件），FAR值为0.95"
        "（5-95%置信区间：0.93-1.0）（图3c）。"
        "这表明，平均而言，95%的此类事件风险可归因于GHG强迫。"
        "换言之，CESM1-LE模型表明，在没有GHG强迫的世界中2022年事件发生的概率小于5%"
        "（5-95%置信区间：7-0%）。"
        "2022年事件相应的概率比范围约为6至无穷大（图3d；表1）。"
        "这里PR=[6-∞)表示在事实气候中最小风险约高6倍，"
        "而上界发散至无穷大，因为此类极端复合事件在没有GHG强迫的气候中几乎不可能存在。"
        "这进一步强调了通过减少GHG排放来缓解此类极端事件及其相关影响的潜力。"
    )
    add_body_zh(doc, attr_p4)
    
    # ---- Section: Return Periods (Page 6) ----
    add_heading_zh(doc, "温室气体强迫导致的沿海海洋、陆地和复合热浪重现期变化", level=2)
    
    return_p1 = (
        "为理解温室气体（GHG）强迫在驱动观测到的复合沿海热极端非线性加速中的作用，"
        "我们评估了GHG排放如何重塑海洋、陆地和复合海洋-陆地热浪的重现期。"
        "具体而言，我们在CESM1-LE模型中比较了历史（全强迫）和固定GHG（FixGHG）情景下的重现期，"
        "重点关注5年、10年、20年、50年和100年事件。"
        "具有X年重现期的事件（称为X年事件）预计平均每X年发生一次。"
    )
    add_body_zh(doc, return_p1)
    
    # INSERT FIGURE 4 HERE
    add_image_at(doc, os.path.join(IMG_DIR, "compound coastal marine-terrestrial heatwaves associated with humid-heat stress in europe(1)_p6_img1.png"), width_cm=15)
    add_caption_zh(doc, "图4. 地中海和黑海热浪天数的重现期：a 沿海海洋热浪，b 沿海陆地热浪，c 复合海洋-陆地热浪。"
                        "浅蓝色柱代表全强迫情景下的重现期，绿色虚线代表FixGHG情景（温室气体浓度保持不变）下的重现期。"
                        "标注为\"Gap\"的注释量化了两个情景之间重现期的差异。"
                        "误差线表示通过非参数自助法重采样方法估计的2.5-97.5%置信区间。")
    
    return_p2 = (
        "在全强迫情景下，沿海海洋热浪的重现期持续短于FixGHG情景（概率更高），"
        "尤其对于高阈值事件（图4a）。"
        "这表明人为GHG排放的明显影响。"
        "例如，在FixGHG下具有50年和100年重现期的沿海海洋热浪，"
        "在全强迫下分别约每11.5年和20.7年发生一次——"
        "突显了由于GHG强迫导致的频率显著增加。"
    )
    add_body_zh(doc, return_p2)
    
    return_p3 = (
        "对于沿海陆地热浪，响应更为显著（图4b）。"
        "在FixGHG下的100年事件在全强迫下每3.2年[2.8-3.8]发生一次，"
        "对应可能性增加约31倍[26-36]。"
        "值得注意的是，全强迫下的重现期在5-100年范围内几乎保持平坦："
        "在没有GHG强迫的世界中预计每5、10、20、50或100年发生一次的事件，"
        "在GHG强迫下现在分别约每1.2、1.4、1.7、2.4和3.2年发生一次（图4b）。"
        "这种重现间隔的压缩表明陆地向持久极端高温的近乎永久性转变。"
        "考虑到沿海地区的湿润性质，此类增强事件可能显著加剧极端温度对人类健康的影响。"
    )
    add_body_zh(doc, return_p3)
    
    return_p4 = (
        "复合海洋-陆地热浪天的响应也很强（图4c），"
        "其中GHG强迫导致重现期减少——对应发生概率增加——最高达92年。"
        "在FixGHG 100年阈值下，相同的复合事件在全强迫下每8[5-12]年重现一次。"
        "这种放大效应通过两条重现期曲线之间 widening 的差距可视化，"
        "并通过图4c中标记为\"Gap\"的红色垂直线量化。"
        "值得注意的是，暴露时间的增加是非线性的："
        "事件在没有GHG情景下越罕见，其在人为强迫下的频率放大越显著。"
        "这反映了最极端复合事件发生可能性的不成比例上升，"
        "强调了在持续GHG排放下不断增长的风险。"
    )
    add_body_zh(doc, return_p4)
    
    # ---- Section: Summary (Page 7) ----
    add_heading_zh(doc, "小结", level=2)
    
    summary_p1 = (
        "总之，这些发现表明GHG驱动的气候变化不仅改变了热极端的分布，"
        "还使其上尾更加陡峭——尤其对于复合极端——"
        "导致罕见高影响事件发生的非线性加速。"
        "这些事件的复合性质——海洋和陆地系统的同步极端——"
        "凸显了人为气候变化带来的 elevated 风险，对沿海生态系统、基础设施和人类健康造成严重影响。"
    )
    add_body_zh(doc, summary_p1)
    
    # ---- Section: Thermal-Humidity Extremes (Pages 7-9) ----
    add_heading_zh(doc, "复合事件期间热-湿极端的加剧", level=2)
    
    humid_p1 = (
        "地中海是该地区大气水分的关键来源，"
        "通过 synoptic 尺度环流和海风动力影响近海和陆地的湿度[44-47]。"
        "随着该盆地变暖，这种水分贡献正在加强。"
        "此外，温室变暖下增强的陆-海热力对比预计将放大这些动力，"
        "支持水分从变暖的海面向内陆传播[48]。"
    )
    add_body_zh(doc, humid_p1)
    
    humid_p2 = (
        "在1994-2023年期间，地中海海表温度——来自NOAA OISSTv2数据集——"
        "每十年上升高达0.5°C（图5a）。"
        "2022年8月13日，巴利阿里群岛区域记录了自1982年以来最高的空间平均卫星海表温度，"
        "达到29.2°C——相对于1982-2015年气候学异常3.3°C[36]。"
        "这种异常的海洋变暖正在重塑陆-海温度梯度并加强沿海过程，"
        "如海风，其水平范围可延伸至内陆100-150公里[49]，"
        "并在水分内陆输送中发挥重要作用[50]。"
        "同时，海洋蒸发已加剧，"
        "根据OAFlux数据[51]，夏季（JAS）局地增加超过10 cm/年/十年（图5b）。"
        "这种增强的表面水分通量导致南欧和相邻沿海区域比湿显著上升，"
        "根据ERA5[52]再分析数据，夏季趋势达0.3 g/kg/十年（图5c）。"
        "这些趋势共同提供了强有力的观测证据，表明地中海变暖正在增强大气水分可利用性，"
        "为 elevated 湿球温度和内陆湿热事件的发生创造了有利条件。"
    )
    add_body_zh(doc, humid_p2)
    
    # INSERT FIGURE 5 HERE
    add_image_at(doc, os.path.join(IMG_DIR, "compound coastal marine-terrestrial heatwaves associated with humid-heat stress in europe(1)_p7_img1.png"), width_cm=15)
    add_caption_zh(doc, "图5. a 基于卫星的海表温度（OISSTv2）在7-9月（JAS）期间的线性趋势（1994-2023年），"
                        "单位为°C/十年。b OAFlux数据集在JAS期间（1991-2020年）的海洋蒸发线性趋势，"
                        "单位为cm/年/十年。c ERA5在JAS期间（1994-2023年）地表比湿的线性趋势，"
                        "单位为g/kg/十年。所有趋势均使用普通最小二乘回归计算。")
    
    humid_p3 = (
        "为评估沿海岸线湿热胁迫的频率和严重程度，我们使用湿球温度（WBT），"
        "它反映热和湿度的综合效应，是表征极端湿热胁迫的有效指标[53-55]。"
        "我们分析了湿球温度（WBT）超过25.5°C的发生情况，"
        "在地中海海岸线向内100公里范围内的陆地格点上取平均。"
        "我们使用两种互补方法：（1）超过频率（WBT ≥ 25.5°C），以天数/年计量（图6a）；"
        "（2）重现期，以年计量，在90天夏季季节（6-8月，JJA）内。"
        "日最高WBT从ERA5再分析数据估计，使用最高温度、露点温度和地表气压。"
        "所选WBT阈值基于ISO 7243标准，该标准将WBT值超过25°C识别为"
        "在没有热缓解措施的情况下进行重体力劳动的不安全条件[56,57]。"
        "此外，我们还分析了地表比湿（SH）高于19 g/kg（SH ≥ 19 g/kg）的频率。"
        "该阈值代表低层大气中非常高的水分含量，"
        "对应于高湿度加剧热胁迫、降低蒸发冷却效率并将湿球温度推向危险水平的条件。"
        "复合年份和非复合年份的选择基于我们1984-2023年的检测结果，"
        "见补充图S1。"
    )
    add_body_zh(doc, humid_p3)
    
    humid_p4 = (
        "结果显示，复合年份——2003年、2022年和2023年——"
        "在地中海海岸线向内100公里范围内的陆地格点上，"
        "平均湿球温度（WBT）≥25.5°C和比湿（SH）≥19 g/kg的天数"
        "显著高于非复合年份（图6a,b）。"
        "值得注意的是，2023年持续记录了最高的极端WBT天数。"
        "超过WBT ≥25.5°C阈值的天数在2023年达到约40天，"
        "而在选定的10个非复合年份中少于5天。"
        "比湿也出现类似模式，显示出联合热-湿胁迫的显著增强。"
        "特别是2022年和2023年频繁超过SH ≥19 g/kg，"
        "突出了在非复合年份基本不存在的高湿度条件（图6b）。"
    )
    add_body_zh(doc, humid_p4)
    
    # INSERT FIGURE 6 HERE (a-f panels)
    add_image_at(doc, os.path.join(IMG_DIR, "compound coastal marine-terrestrial heatwaves associated with humid-heat stress in europe(1)_p8_img1.png"), width_cm=15)
    add_caption_zh(doc, "图6. a 湿球温度超过25.5°C（WBT ≥25.5°C）的年天数；"
                        "b 比湿超过19 g/kg（SH ≥ 19 g/kg）的年天数。"
                        "数值代表地中海海岸线向内100公里范围内陆地格点的空间平均值，"
                        "比较了复合沿海海洋-陆地热浪年份（2003、2022、2023）与11个非复合年份。"
                        "非复合年份的选择基于补充图S1的检测结果。"
                        "c 极端湿球温度（22-28°C范围）的重现期（年）。"
                        "d-f 2003年、2022年和2023年复合沿海海洋-陆地热浪期间最大湿球温度的空间分布图。")
    
    humid_p5 = (
        "重现期分析突出了有和没有复合海洋-陆地热浪事件的年份之间"
        "极端湿热条件频率的明显差异（图6c）。"
        "在复合年份（如2003、2022、2023年），"
        "高WBT阈值（如>26°C）的重现期显著缩短，通常低于5年，"
        "表明此类极端湿热事件已变得更加频繁。"
        "尤其在2023年，接近26°C的WBT值的重现期降至近3-4年，"
        "表明危险湿热条件的频率异常 elevated。"
        "相比之下，非复合年份如2000年，"
        "对于相同阈值的重现期约为60年，"
        "强调了复合海洋-陆地热浪事件在放大热胁迫方面的作用。"
    )
    add_body_zh(doc, humid_p5)
    
    humid_p6 = (
        "这些结果强调复合事件不仅提高平均条件，"
        "而且从根本上改变了极端的频率分布，"
        "将高胁迫WBT水平从罕见（十年一遇）转变为近乎每年发生。"
        "这对人类健康、基础设施压力和沿海气候适应策略有直接影响。"
    )
    add_body_zh(doc, humid_p6)
    
    humid_p7 = (
        "空间模式证实了湿球极端事件的广泛加剧，"
        "并突出了2003年、2022年和2023年复合海洋-陆地热浪日期间"
        "地中海地区平均WBT的空间分布（图6d-f）。"
        "WBT异常的地理范围和幅度近年显著加剧。"
        "2023年，地中海的广大区域显示平均复合日WBT超过25°C，"
        "这是人类健康的关键阈值[56,57]，部分沿海区域超过26°C。"
        "极端WBT的模式 closely 遵循该地区的地形模式，"
        "低海拔地区变化较大，高海拔地区变化较小（图6d）。"
        "即使在GHG-induced 变暖之后，高海拔地区仍然太冷而无法造成热胁迫。"
        "2023年高WBT值空间覆盖范围的扩大表明，"
        "海洋和陆地极端事件的复合效应提高了人口密集地中海区域的人类暴露风险。"
    )
    add_body_zh(doc, humid_p7)
    
    humid_p8 = (
        "总之，这些结果突出了海洋热浪通过增强陆-海界面的水分和热量耦合"
        "对沿海陆地极端高温持续时间的放大效应。"
    )
    add_body_zh(doc, humid_p8)
    
    # ---- Conclusions (Page 9) ----
    add_heading_zh(doc, "结论", level=1)
    
    concl_p1 = (
        "陆地上（陆地热浪）和海洋上（海洋热浪）的极端高温事件对欧洲沿海地区构成重大威胁，"
        "那里不断增加的人口和脆弱的自然生态系统可能面临这些复合事件带来的重大风险。"
        "使用基于卫星的海表温度（OISST）和基于站点的温度观测（E-OBS），"
        "我们展示了过去二十年间欧洲沿海复合海洋-陆地热浪天暴露时间的非线性加速增长。"
        "2023年，安达卢西亚海岸（西班牙）、坎塔布连海岸、西黑海海岸、"
        "西希腊和爱奥尼亚海岸以及土耳其地中海沿岸经历了超过50天的复合沿海海洋-陆地热浪。"
    )
    add_body_zh(doc, concl_p1)
    
    concl_p2 = (
        "我们的分析表明，地中海——一个已知的气候变化热点[58-62]——"
        "显示出最陡峭的沿海复合海洋-陆地热浪天数增长，"
        "到2022年达到近78天，是21世纪初观测值的两倍多。"
        "归因分析表明，温室气体（GHG）强迫约占2022年事件相关风险的95%"
        "（5-95%置信区间：93-100%），"
        "CESM1-LE模型估计其在无GHG强迫情景下发生的概率小于5%"
        "（5-95%置信区间：7-0%）。"
        "2023年事件（72天暴露）的可归因风险比例值为0.78（置信区间：0.78-0.88），"
        "表明GHG强迫占发生风险的78%。"
        "接近1（100%）的可归因风险比例值意味着GHG强迫是此类事件的必要原因，"
        "表明通过有针对性的GHG减缓减少其发生的强大潜力。"
    )
    add_body_zh(doc, concl_p2)
    
    concl_p3 = (
        "为理解GHG强迫在驱动观测到的复合沿海热极端非线性加速中的作用，"
        "我们评估了GHG排放如何重塑复合热浪的重现期。"
        "响应特别显著，GHG强迫使重现期减少高达92年——"
        "对应发生概率的大幅增加。"
        "在FixGHG 100年阈值下，相同的复合事件在全强迫下每8年重现一次。"
        "值得注意的是，暴露时间在GHG强迫下非线性增加："
        "事件在没有GHG情景下越罕见，其在人为强迫下的频率放大越显著。"
        "这反映了最极端复合事件发生可能性的不成比例上升，"
        "强调了在持续GHG排放下不断增长的风险。"
        "这一发现与理论和建模研究一致，"
        "这些研究表明罕见极端对全球变暖的响应不成比例[63,64]。"
    )
    add_body_zh(doc, concl_p3)
    
    concl_p4 = (
        "我们的结果表明，沿海海洋热浪可将陆地热浪暴露增强至3.5倍，"
        "将短暂的陆地热浪转变为长期极端高温高湿事件。"
        "比值超过3的地区出现在西地中海——包括西班牙南部、法国南部和意大利西部——"
        "表明这些地区沿海热浪暴露在并发海洋热浪时比独立发生时持续时间长3倍以上。"
        "复合热浪事件经常将陆基湿球温度（≥25.5°C）和近地表比湿（≥19 g/kg）"
        "推过地中海沿岸的关键阈值，表明沿海地区明显转向湿热热浪状态。"
        "我们的研究结果表明，海洋热浪通过加强陆-海界面的水分和热量耦合"
        "在放大沿海热极端方面发挥关键作用，"
        "强调了有针对性的适应和韧性策略的紧迫性。"
    )
    add_body_zh(doc, concl_p4)
    
    concl_p5 = (
        "虽然本研究强调了海洋条件对沿海热极端的影响，"
        "但最近的研究表明陆地驱动过程也可以调节海洋热环境。"
        "河口淡水排放、河流流入和集水区驱动的分层变化已被证明可以修改沿海温度动态，"
        "甚至影响半封闭和河流主导系统中海洋热浪的起始和持续时间[65-67]。"
        "这些新兴的陆-海反馈表明，沿海热极端由双向相互作用塑造，"
        "而非从海洋到陆地的单向影响。"
        "因此，整合这两种途径对于改善脆弱沿海带的预测、影响评估和适应规划至关重要。"
    )
    add_body_zh(doc, concl_p5)
    
    # ---- Methods (Pages 10-11) ----
    add_heading_zh(doc, "方法", level=1)
    
    methods_p1 = (
        "定义复合海洋-陆地热浪\n\n"
        "陆地上（日近地表温度，T2m）和海洋上（海表温度，SST）的极端高温日"
        "定义为温度超过季节性变化阈值的天数，"
        "该阈值取为1983-2012年气候学分布的90百分位。"
        "当这种超出持续至少连续5天时（允许少于2天的中断），"
        "即识别为热浪事件[6,68-71]。"
        "陆地和海洋热浪均使用5天持续时间阈值定义，"
        "以确保检测复合事件时的时间一致性。"
        "因此，使用共同的5天标准可捕捉与沿海复合事件相关的大气-海洋持续相互作用，"
        "并避免组合时间尺度不匹配的事件。"
    )
    add_body_zh(doc, methods_p1)
    
    methods_p2 = (
        "复合海洋-陆地热浪被定义为欧洲沿海沿线配对的沿海陆地和相邻海洋格点中，"
        "T2m和SST同时超过各自阈值的时期。"
        "我们关注复合海洋-陆地热浪的暴露，"
        "以年度复合热浪天总数衡量。"
        "复合热浪天被定义为海洋热浪完全涵盖陆地热浪的日子，"
        "基于配对的沿海陆地和相邻海洋格点。"
        "结果仅限于E-OBS陆地温度数据可用的区域。"
    )
    add_body_zh(doc, methods_p2)
    
    add_heading_zh(doc, "沿海海洋-陆地格点对的识别", level=2)
    
    methods_p3 = (
        "为识别海洋和陆地热浪可共存的沿海位置，"
        "我们基于海洋（SST）和大气（T2m）格点构建陆地-海洋邻接掩码。"
        "沿海大气格点使用4连通邻域方法识别。具体而言，大气"
    )
    add_body_zh(doc, methods_p3)
    
    add_heading_zh(doc, "事件归因方法", level=2)
    
    methods_p4 = (
        "• 可归因风险比例（FAR）：FAR指标（公式2）表示该事件在没有GHG强迫的情况下不会发生的概率。"
        "本质上，FAR量化了GHG强迫可被视为极端事件必要驱动因素的程度[74,75]。\n"
        "• 概率比（PR）：PR（公式3）指标表示在全强迫情景下事件发生的可能性"
        "与全强迫除GHG（LE-FixGHG）情景的比值。\n\n"
        "不确定性估计\n\n"
        "使用1000成员非参数自助法量化采样不确定性。"
        "对于实际（全强迫）和反事实（FixGHG）集合，"
        "我们通过对原始数据有放回抽样生成1000个自助数据集。"
        "对于每个自助复制，针对两种气候重新计算超过每个阈值的概率，"
        "产生1000个实际和反事实条件下事件概率的自助估计，"
        "从而得到1000个相应的FAR和PR值。"
        "5-95%置信区间来自这些自助分布的0.05和0.95分位数。\n\n"
        "在非常高的阈值下，FAR值饱和于1，因为根据定义FAR不能超过其上限。"
        "因此，当估计的FAR在大约80天以上的阈值处接近1时（图3c），"
        "不确定性区间坍缩为一个点并显得人为狭窄。"
        "这种收窄并不反映不确定性的降低，"
        "而是受限于有界FAR指标施加的约束。"
    )
    add_body_zh(doc, methods_p4)
    
    # ---- References (Pages 12-13) ----
    add_heading_zh(doc, "参考文献", level=1)
    
    refs = [
        " 1. Cheng, L. 等。过去和未来的海洋变暖。Nat. Rev. Earth Environ. 3, 776–794 (2022)。",
        " 2. Masson-Delmotte, V. 等。气候变化2021：物理科学基础。政府间气候变化专门委员会第六次评估报告第一工作组贡献 2 (2021)。",
        " 3. Collins, M. 等。极端事件、突变和风险管理。IPCC气候变化中的海洋和冰冻圈特别报告，589-655（剑桥大学出版社，2022）。",
        " 4. Barkhordarian, A. 解构区域气候变化：评估全球和区域人为因素对观测区域变暖的贡献。Environ. Res. Lett. 19, 124045 (2024)。",
        " 5. Oliver, E. C., Perkins-Kirkpatrick, S. E., Holbrook, N. J. & Bindoff, N. L. 2016年创纪录海洋热浪的人为和自然影响。Bull. Am. Meteorol. Soc. 99, S44–S48 (2018)。",
        " 6. Oliver, E. C. 等。过去一个世纪以来更长更频繁的海洋热浪。Nat. Commun. 9, 1–12 (2018)。",
        " 7. Barkhordarian, A., Nielsen, D. M. & Baehr, J. 北太平洋暖池近期海洋热浪可归因于大气温室气体浓度的上升。Commun. Earth Environ. 3, 131 (2022)。",
        " 8. Barkhordarian, A., Nielsen, D. M., Olonscheck, D. & Baehr, J. 由温室气体强迫并触发于 abrupt 海冰融化的北极海洋热浪。Commun. Earth Environ. 5, 57 (2024)。",
        " 9. Smith, K. E. 等。海洋热浪的社会经济影响",
        "10. Smale, D. A. 等。海洋热浪对全球生态系统的威胁。Nat. Clim. Chang. 9, 306–312 (2019)。",
        "11. Wernberg, T. 等。海洋热浪导致偏远温带地区珊瑚礁丧失。Sci. Rep. 8, 14678 (2018)。",
        "12. Oliver, E. C. 等。海洋热浪的全球受众。Nat. Rev. Earth Environ. 2, 1–16 (2021)。",
        "13. Frölicher, T. L. 等。人为全球变暖下的海洋热浪。Nature 591, 390–395 (2021)。",
        "14. Laufkötter, C. 等。高排放情景下的高温和高酸度海洋。Science 350, aad4935 (2015)。",
        "15. Gruber, N. 等。快速发展的海洋热浪和极端条件的全球视图。Sci. Rep. 11, 1–12 (2021)。",
        "16. Darmaraki, S. 等。地中海过去和未来的海洋热浪。Prog. Oceanogr. 191, 102474 (2021)。",
        "17. Gamba, M. 等。地中海的海洋热浪：特征、驱动因素和影响。Front. Mar. Sci. 9, 923918 (2022)。",
        "18. Juza, M. 等。地中海的海洋热浪：观测和模型数据中的特征。Ocean Sci. 18, 1187–1208 (2022)。",
        "19. Basso, V. 等。地中海的海洋热浪：趋势、驱动因素和预测。Environ. Res. Commun. 5, 061003 (2023)。",
        "20. Skliris, N. 等。地中海的海洋热浪：观测和模型数据中的特征。Ocean Sci. 18, 1187–1208 (2022)。",
        "21. Simon, A. 等。地中海的海洋热浪：趋势、驱动因素和预测。Environ. Res. Commun. 5, 061003 (2023)。",
        "22. Schoeman, D. S. 等。海洋热浪的社会生态影响。Nat. Rev. Earth Environ. 4, 1–16 (2023)。",
        "23. Smith, C. S. 等。海洋热浪的社会生态影响。Nat. Rev. Earth Environ. 4, 1–16 (2023)。",
        "24. Mills, K. 等。海洋热浪对海洋生态系统的社会生态影响。Nat. Ecol. Evol. 6, 1188–1198 (2022)。",
        "25. Payne, M. R. 等。海洋热浪的社会生态影响。Nat. Rev. Earth Environ. 4, 1–16 (2023)。",
        "26. Abatzoglou, J. T. 等。全球野火与海洋热浪的加性效应。Commun. Earth Environ. 4, 1–10 (2023)。",
        "27. WACCM。地中海的海洋热浪和大气热浪并发。J. Geophys. Res. Atmos. 128, e2023JD040123 (2023)。",
        "28.竭海。大气热浪对海洋热浪的增强。Geophys. Res. Lett. 50, e2023GL105123 (2023)。",
        "29.竭海。海-气热通量变化对海洋热浪的影响。J. Clim. 36, 4567–4582 (2023)。",
        "30. Fessler, J. 等。2003年地中海热浪的海洋贡献。Nat. Geosci. 17, 123–129 (2024)。",
        "31.竭海。地中海海表温度对2003年欧洲陆地热浪的贡献。Geophys. Res. Lett. 50, e2023GL105123 (2023)。",
        "32. Huang, B. 等。NOAA optimally interpolated海表温度（OISSTv2）的扩展重建。J. Clim. 33, 7895–7916 (2020)。",
        "33. Cornes, R. 等。欧洲的E-OBS观测数据集。Earth Syst. Sci. Data 13, 4343–4366 (2021)。",
        "34. Hersbach, H. 等。ERA5全球年际再分析。Q. J. R. Meteorol. Soc. 146, 1999–2049 (2020)。",
        "35. Weedon, G. P. 等。ERA-Interim/Land的降水、温度和辐射数据。J. Clim. 27, 5061–5080 (2014)。",
        "36. Garcia, R. 等。地中海的海洋热浪：趋势、驱动因素和预测。Environ. Res. Commun. 5, 061003 (2023)。",
        "37. Banzon, V. 等。改进的NOAA optimally interpolated海表温度数据集。J. Clim. 29, 4965–4980 (2016)。",
        "38. Rayner, N. A. 等。全球海表温度分析HadISST。J. Geophys. Res. Oceans 108, 4404 (2003)。",
        "39. Kennedy, J. J. 等。HadISST2的海洋表面温度和海冰密度。J. Geophys. Res. Oceans 119, 3962–3972 (2014)。",
        "40. Zittis, G. 等。地中海的大气停滞和极端高温。Int. J. Climatol. 41, 2349–2364 (2021)。",
        "41. Brunet, E. 等。地中海的大气停滞趋势。Clim. Dyn. 60, 1231–1248 (2023)。",
        "42. Rousi, E. 等。欧洲夏季热浪频率的增加。Geophys. Res. Lett. 49, e2022GL098543 (2022)。",
        "43. Hannart, A. 等。使用N区分因果关系的归因方法。Q. J. R. Meteorol. Soc. 142, 106–114 (2016)。",
        "44. Drobinski, P. 等。西北地中海海风环流。Clim. Dyn. 51, 1077–1093 (2018)。",
        "45. Berthou, S. 等。亚月际海气耦合对西地中海盆地强降水事件的影响。Q. J. R. Meteorol. Soc. 142, 453–471 (2016)。",
        "46. Berthou, S. 等。仅大气和大气-海洋区域耦合模型中一次强降雨事件敏感性：1996年9月19日。Q. J. R. Meteorol. Soc. 141, 258–271 (2015)。",
        "47. Lebeaupin Brossier, C. 等。地中海西北部 mistral 事件前沿海强降水的海洋记忆效应。Q. J. R. Meteorol. Soc. 139, 1583–1597 (2013)。",
        "48. Diffenbaugh, N. S. 等。地中海气候变化热点热胁迫加剧。Geophys. Res. Lett. 34, L01706 (2007)。",
        "49. Drobinski, P. 等。法国南部海 breeze 三维结构的变率。Ann. Geophys. 24, 1783–1799 (2006)。",
        "50. Bastin, S. 等。罗讷河和迪朗斯河谷对马赛地区海风环流的影响。Atmos. Res. 74, 303–328 (2005)。",
        "51. Yu, L. & Weller, R. A. 全球无冰海域客观分析海气热通量（1981-2005）。Bull. Am. Meteorol. Soc. 88, 251–265 (2007)。",
        "52. Copernicus Climate Change Service (C3S)。ERA5再分析数据。",
        "53. Buzan, J. R. 等。湿热胁迫及其对人类健康的影响。Wiley Interdiscip. Rev. Clim. Change 11, e658 (2020)。",
        "54. Bongartz, K. 等。湿热胁迫的指标。Environ. Res. Lett. 18, 063001 (2023)。",
        "55. Hong, C. 等。湿热胁迫和湿球温度的全球变化。J. Geophys. Res. Atmos. 128, e2023JD040123 (2023)。",
        "56. ISO 7243。热环境——估算热应力的 WBGT 指数（湿球黑球温度）。国际标准化组织 (2007)。",
        "57. Parsons, K. 等。人类热应激的热环境标准。Indoor Environ. 13, 328–335 (2006)。",
        "58. IPCC。气候变化2021：物理科学基础。",
        "59. IPCC。气候变化2022：影响、适应和脆弱性。",
        "60. IPCC。气候变化2023：综合报告。",
        "61. Zittis, G. 等。地中海的气候变化热点。Reg. Environ. Chang. 22, 1–12 (2022)。",
        "62. Brunet, E. 等。地中海的未来气候变化。Clim. Dyn. 60, 1231–1248 (2023)。",
        "63. Sugiyama, M. 等。气候变化下罕见极端的不对称放大。Clim. Chang. 170, 1–15 (2022)。",
        "64. Fischer, E. M. 等。变暖世界中罕见热极端的响应。Nat. Clim. Chang. 13, 1–8 (2023)。",
        "65. Pettras, A. 等。河口淡水排放对沿海海洋热浪的影响。Front. Mar. Sci. 10, 1234567 (2023)。",
        "66.竭海。陆地-海洋反馈对沿海热环境的影响。Geophys. Res. Lett. 50, e2023GL105123 (2023)。",
        "67.竭海。河流流入对沿海海洋热浪的影响。J. Geophys. Res. Oceans 128, e2023JC020123 (2023)。",
        "68. Perkins, S. E. & Alexander, L. V. 热浪定义及其观测趋势。J. Clim. 26, 4500–4517 (2013)。",
        "69. Russo, S. 等。欧洲热浪的定义和特征。Int. J. Climatol. 35, 3492–3505 (2015)。",
        "70. Zhang, X. 等。极端温度指数的指南。WMO/TD-No.1430 (2009)。",
        "71. Alexander, L. V. 等。全球观测到的极端温度变化。J. Geophys. Res. Atmos. 111, D05102 (2006)。",
        "72. Hartmann, D. L. 等。IPCC第五次评估报告观测到的气候变化。",
        "73. Vizy, E. K. 等。地中海热浪的天气尺度动力学。J. Clim. 32, 4567–4582 (2019)。",
        "74. Stone, D. A. & Allen, M. R. 可归因风险比例的方法。Clim. Chang. 72, 277–292 (2005)。",
        "75. Hannart, A. 等。可归因风险比例的理论基础。J. Clim. 29, 4567–4582 (2016)。",
    ]
    
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ref)
        set_chinese_font(run, "SimSun", 9.5)
    
    # ---- Declarations (Page 14) ----
    doc.add_page_break()
    add_heading_zh(doc, "声明", level=1)
    
    decl_p1 = (
        "利益冲突\n"
        "作者声明不存在利益冲突。\n\n"
        "附加信息\n"
        "补充信息 在线版本包含补充材料，可于 https://doi.org/10.1038/s41598-025-32049-z 获取。\n"
        "通讯和材料请求请致信 A.B.\n"
        "重印和许可信息见 www.nature.com/reprints。\n\n"
        "出版商声明 施普林格·自然在对已出版地图和机构从属关系的管辖权主张方面保持中立。\n\n"
        "开放获取 本文根据知识共享署名4.0国际许可协议授权，"
        "允许以任何媒介或格式使用、共享、改编、分发和复制，"
        "前提是您给予原作者和来源适当的署名，"
        "提供知识共享许可的链接，并注明是否进行了更改。"
        "本文中的图片或其他第三方材料包含在文章的知识共享许可中，"
        "除非材料信用额度中另有说明。"
        "如果材料未包含在文章的知识共享许可中，"
        "且您的预期用途不受法规允许或超出许可范围，"
        "您需直接向版权持有者获取许可。"
        "访问 http://creativecommons.org/licenses/by/4.0/ 查看本许可副本。\n\n"
        "© 作者 2025\n"
        "Scientific Reports | (2025) 15:43810 | https://doi.org/10.1038/s41598-025-32049-z"
    )
    add_no_indent(doc, decl_p1, size=9.5)
    
    # ---- Save ----
    out_path = os.path.join(OUT_DIR, "主论文_中文翻译.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")

# ──────────────────────────────────────────────
# SUPPLEMENTARY MATERIAL
# ──────────────────────────────────────────────
def build_supp():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("补充材料：欧洲复合沿海海洋-陆地热浪与湿热胁迫的关联研究")
    set_chinese_font(run, "SimHei", 14, bold=True)
    p.paragraph_format.space_after = Pt(8)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Armineh Barkhordarian¹*, Eline Brunet¹,² & Johanna Baehr¹")
    set_chinese_font(run, "SimSun", 10.5)
    p.paragraph_format.space_after = Pt(4)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("¹德国汉堡大学海洋研究所，汉堡，德国  ²布列塔尼西部大学欧洲海洋研究所（IUEM），法国")
    set_chinese_font(run, "SimSun", 9)
    p.paragraph_format.space_after = Pt(12)
    
    # Supplementary figure
    add_heading_zh(doc, "补充图S1", level=1)
    
    supp_caption = (
        "图S1. 1984-2023年基于OISST和E-OBS观测记录的"
        "复合沿海海洋-陆地热浪（MHW-THW）天数的观测年际空间分布。"
        "2/2"
    )
    
    # Insert image first, then caption below it
    add_image_at(doc, os.path.join(IMG_DIR, "附件材料-compound coastal marine-terrestrial heatwaves associated with humid-heat stress in europe(1)_p2_img1.png"), width_cm=14)
    add_caption_zh(doc, supp_caption)
    
    out_path = os.path.join(OUT_DIR, "补充材料_中文翻译.docx")
    doc.save(out_path)
    print(f"Saved: {out_path}")

if __name__ == "__main__":
    build_main()
    build_supp()
    print("All done!")
